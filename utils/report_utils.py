# --- utils/report_utils.py ---
import io
import csv
from docx import Document
from docx.shared import Inches

def generate_csv(data_list, fieldnames):
    """
    Génère un contenu CSV en mémoire.

    Args:
        data_list (list): Liste de dictionnaires.
        fieldnames (list): Noms des colonnes.

    Returns:
        str: Contenu CSV sous forme de chaîne.
    """
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for row in data_list:
        writer.writerow({k: row.get(k, "") for k in fieldnames})
    return output.getvalue()

def generate_docx_report(data_dict, title="Rapport"): 
    """
    Génère un fichier DOCX simple à partir d'un dictionnaire de données.

    Args:
        data_dict (dict): Données clef-valeur.
        title (str): Titre du document.

    Returns:
        BytesIO: Flux binaire du fichier DOCX.
    """
    document = Document()
    document.add_heading(title, level=1)

    for key, value in data_dict.items():
        document.add_paragraph(f"{key}: {value}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
