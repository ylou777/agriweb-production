"""
Génère le document Word de réponse à NEWS-SOLAR (proposition de collaboration logicielle)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ───────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(30, 90, 160))
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(50, 120, 50))
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def shade_cell(cell, hex_color="D6E4F0"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "1F5C9E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
    # Data rows
    for r_idx, row in enumerate(rows):
        shade = "EAF3FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            shade_cell(cell, shade)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_font(run, size=10)
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════════════════════
#  EN-TÊTE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f"Le {datetime.date.today().strftime('%d/%m/%Y')}")
set_font(run, size=10, italic=True, color=(120, 120, 120))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROPOSITION DE COLLABORATION LOGICIELLE")
set_font(run, size=18, bold=True, color=(20, 70, 140))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Logiciel de simulation et de gestion — Technologie HST NEWS-SOLAR")
set_font(run, size=13, italic=True, color=(60, 60, 60))

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Objet : ")
set_font(run, size=11, bold=True)
run2 = p.add_run("Réponse détaillée à votre proposition de co-édition d'un logiciel adapté à la technologie HST NEWS-SOLAR")
set_font(run2, size=11)
p.paragraph_format.space_after = Pt(6)

body(doc, "Monsieur,")
body(doc, (
    "Je vous remercie pour votre message détaillé et la clarté de votre vision. Ayant pris connaissance "
    "de l'ensemble de vos pages techniques (Technologie, Décarbonation Complète, Convertisseurs, "
    "Produits), je suis en mesure de vous répondre précisément sur chacun des points fonctionnels "
    "évoqués, ainsi que sur le cadre de notre collaboration."
))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE I — ANALYSE POINT PAR POINT
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "I. Analyse fonctionnelle point par point")

items = [
    (
        "1. Captation solaire pure/brute selon irradiance / GPS",
        "Notre plateforme intègre déjà un moteur de calcul d'irradiance géolocalisé (API Copernicus / PVGIS / "
        "Solargis selon la région). Il s'agit d'étendre ce moteur pour y intégrer le coefficient de captation "
        "spécifique au procédé HST (95 % vs ~20 % PV), le ratio de superficie active (1/10 000 de la surface de "
        "réception), ainsi que les données constructeur de vos concentrateurs (9,5 MWc thermiques/ha). "
        "Techniquement réalisable dès la phase bêta."
    ),
    (
        "2. Quantité d'énergie pure/brute disponible pour des coordonnées GPS données",
        "Module de simulation énergétique par coordonnées GPS : calcul du GHI (Global Horizontal Irradiance), "
        "DNI (Direct Normal Irradiance) — particulièrement pertinent pour votre procédé à hyper-concentration — "
        "et intégration du profil climatologique annuel. Restitution sous forme de MWh/an disponibles, avec "
        "profil saisonnier et horaire. Directement mappable sur votre technologie."
    ),
    (
        "3. Production estimative : chaleur, froid, électricité, H₂, NH₃, e-fuels",
        "Cœur du moteur de simulation propre NEWS-SOLAR. Il faudra modéliser : le flux HST (champ solaire → "
        "absorbeur → batterie thermique → convertisseur), les rendements de conversion selon la sortie souhaitée "
        "(thermique direct, électrique 35/60 %, H₂/NH₃ via électrolyse HTe, SAF/e-SAF), les équations "
        "thermodynamiques du procédé (Q = m.cp.ΔT, Carnot adapté aux cycles supercritiques CO₂). Ce module est "
        "le plus stratégique et devra s'appuyer sur vos bases de données propriétaires. Il constituera l'essentiel "
        "de la PI partagée."
    ),
    (
        "4. Prise en compte des différentes versions (rendement électro-solaire 35 % et 60 %)",
        "Paramétrage du type de convertisseur : mono-étagé (35 %), bi-étagé (60 %), PhotoStatique multi-jonctions "
        "(>40 %, 400 KWc/m²), cycle de Rankine HTe (>1 MWc). Interface de sélection du convertisseur avec "
        "recalcul automatique des sorties énergétiques et du ROI. Directement intégrable dans le formulaire "
        "de simulation."
    ),
    (
        "5. Gestion par IA propriétaire (brevet NEWS-SOLAR)",
        "Notre architecture actuelle comporte déjà un moteur IA. L'adaptation NEWS-SOLAR implique : un modèle "
        "prédictif d'irradiance et de gestion thermique (charge/décharge batterie), des algorithmes d'optimisation "
        "en temps réel des paramètres de configuration, un module de maintenance prédictive basé sur vos données "
        "MTBF (~220 000 h), un dashboard de performance avec alertes et recommandations automatiques. Vos brevets "
        "devront être clairement délimités contractuellement pour encadrer l'intégration."
    ),
    (
        "6. Comparatif NEWS-SOLAR vs PV et autres dispositifs",
        "Module de benchmarking à adapter avec vos coefficients officiels : électrique ×3, cogénération ×8, "
        "hydrogène ×4, stockage ×1 500. Génération automatique de tableaux comparatifs et graphiques "
        "(CAPEX, production 25 ans, CO₂ évité, ROI). Très fort levier marketing pour vos investisseurs."
    ),
    (
        "7. Ajout et dimensionnement/prix d'une batterie thermique",
        "Moteur de dimensionnement : capacité souhaitée en MWh → calcul du volume (1,3–1,5 MWh/m³), coût estimatif, "
        "intégration possible d'électricité verte externe via résistance intégrée, pertes thermiques (1 %/j), "
        "durée d'autonomie. Module \"batterie mobile VHT\" pour configurations déportées."
    ),
    (
        "8. Trading énergétique chaud/froid/électricité et autres",
        "Module de gestion de l'arbitrage énergétique : intégration des données de marché (EPEX Spot, prix spot "
        "gaz, tarifs réglementés), optimisation de la revente aux heures de pointe, simulation des revenus PPA "
        "(5–25 ans). Connecteurs API vers les marchés européens de l'énergie. Modélisation des crédits carbone "
        "(horizon 100 €/T CO₂)."
    ),
    (
        "9. Spécificités d'installation (sol/toiture/ombrières) — gestion poly-énergies",
        "Paramétrage des typologies d'installation : toiture/ombrière (emprise nulle ou limitée), terrain "
        "proximal, configuration mobile (batterie VHT). Configuration poly-énergies simultanées : chaleur + "
        "froid + électricité + H₂ en sortie parallèle ou séquentielle selon les besoins process du client. "
        "Module de contraintes réglementaires par pays."
    ),
    (
        "10. Suivi distant (datas constructeur sur centrale) de la production",
        "Interface de monitoring IoT/SCADA : ingestion des données constructeur via API REST ou MQTT, affichage "
        "temps réel sur dashboard, historique de production, alertes de dépassement de seuil. Compatible "
        "protocoles industriels standards (Modbus TCP, OPC-UA). Module d'export vers les organismes tiers "
        "(certification, audit)."
    ),
    (
        "11. CAPEX / OPEX et ROI estimés selon paramètres",
        "Module financier complet : calcul automatique du CAPEX à partir de la puissance configurée et du type "
        "d'installation, OPEX (quasi-nul maintenance pour versions hermétiques), ROI sur 25 ans avec ou sans PPA, "
        "calcul des aides d'État (par pays et région), revenus de revente crédits carbone. Génération automatique "
        "de fiches financières synthétiques."
    ),
    (
        "12. Mise à jour logicielle / paramètres selon évolutions matérielles",
        "Architecture modulaire avec table de paramètres constructeur versionnée. Mise à jour des coefficients "
        "techniques sans recompilation. Gestion des versions hardware (concentrateur génération 1/2/N) avec "
        "historique d'impact sur les simulations existantes. Administration sécurisée réservée à NEWS-SOLAR."
    ),
    (
        "13. Déploiement multi-langues et multi-pays avec acquisition de données spécifiques",
        "La plateforme est déjà architecturée pour un déploiement international. Adaptation : internationalisation "
        "i18n complète (multi-langues), acquisition de données d'irradiance locales par pays/région (APIs "
        "spécialisées, données satellitaires), intégration des réglementations énergétiques et fiscales locales, "
        "adaptation des devises et unités."
    ),
    (
        "14. Version Grands Comptes « sensible » — projets d'envergure ou stratégiques",
        "Environnement isolé (tenant dédié ou instance séparée), chiffrement renforcé des données projet, "
        "contrôle d'accès RBAC granulaire, audit trail complet, NDA/confidentialité technique embarqués dans "
        "les flux de validation. Alignement avec ISO 27001 (déjà audité sur notre plateforme existante)."
    ),
    (
        "15. Édition d'une proposition chiffrée estimative (devis simplifié) avant approbation commerciale",
        "Module de génération de devis PDF structuré : récapitulatif de la configuration simulée, chiffres clés "
        "(production, ROI, CO₂), tarification estimative, watermark \"avant approbation commerciale/direction\". "
        "Workflow de validation avec approbation sécurisée côté service commercial NEWS-SOLAR."
    ),
    (
        "16. Cybersécurité afférente",
        "Base existante auditée (OWASP Top 10, authentification multi-facteurs, chiffrement TLS/AES-256, "
        "protection injection, gestion des sessions). À renforcer pour la version Grands Comptes : audit de "
        "pénétration dédié, conformité RGPD/NIS2, gestion des secrets via vault, journalisation sécurisée."
    ),
]

for title, desc in items:
    heading2(doc, title)
    body(doc, desc)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE II — MODELE CO-EDITION
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "II. Modèle de co-édition / partenariat")

body(doc, (
    "Votre proposition de co-éditeur est la seule qui ait du sens sur un projet de cette nature, "
    "pour plusieurs raisons :"
))

bullet(doc, (
    "Propriété intellectuelle : vos bases de données techniques, vos modèles de rendement et vos "
    "algorithmes IA constituent le cœur de valeur. Leur intégration dans le logiciel ne peut se faire "
    "que dans un cadre contractuel clair définissant la co-titularité du logiciel résultant, distincte "
    "de vos brevets hardware/procédé."
))
bullet(doc, (
    "Allocation des ressources : NEWS-SOLAR apporte ses données, expertises métier et validation "
    "technique ; notre équipe apporte l'architecture logicielle, le développement et le déploiement. "
    "Un accord de co-développement (type JDA – Joint Development Agreement) est le cadre juridique approprié."
))
bullet(doc, (
    "Évolutions : la table de paramètres versionnée garantit que chaque évolution matérielle de vos "
    "centrales se traduit par une simple mise à jour de paramètres, sans refonte logicielle."
))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE III — PLAN DE TRAVAIL
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "III. Proposition de plan de travail conjoint")

add_table(
    doc,
    headers=["Phase", "Contenu", "Livrable"],
    rows=[
        ["Phase 0 — Cadrage (2 semaines)",
         "Contractualisation JDA, inventaire des bases PI NEWS-SOLAR, définition du périmètre bêta",
         "Contrat signé, cahier des charges bêta"],
        ["Phase 1 — Bêta simplifiée (6-8 semaines)",
         "Modules : GPS/irradiance, simulation multi-énergies (3 configurations), comparatif PV, CAPEX/ROI, export PDF devis",
         "Version bêta démontrable"],
        ["Phase 2 — Ajustements & prise en main (2-3 semaines)",
         "Sessions de simulation conjointes, recalibrage des modèles, corrections UX",
         "Version bêta validée"],
        ["Phase 3 — Vidéo & présentation investisseurs (2 semaines)",
         "Scénario de simulation live, captures d'écran annotées, montage vidéo",
         "Vidéo + slides investisseurs"],
        ["Phase 4 — Version investisseurs (4-6 semaines)",
         "Modules trading, monitoring IoT, multi-langues (FR/EN), version Grands Comptes",
         "Version pre-production investisseurs"],
        ["Phase 5 — Production & déploiement",
         "Intégration IA propriétaire, mises à jour paramétriques, déploiement international, cybersécurité avancée",
         "Version commerciale V1"],
    ],
    col_widths=[5.5, 8.0, 5.5]
)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE IV — ESTIMATION TARIFAIRE
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "IV. Estimation tarifaire de la prestation")

heading2(doc, "A. Volumétrie de développement par phase")

add_table(
    doc,
    headers=["Phase", "Contenu principal", "Jours estimés"],
    rows=[
        ["Phase 0 — Cadrage & contrat", "Spécifications, architecture, JDA", "3 – 5 j"],
        ["Phase 1 — Bêta simplifiée", "Simulation GPS/énergie, comparatif PV, CAPEX/ROI, devis PDF", "35 – 45 j"],
        ["Phase 2 — Ajustements", "Recalibrage modèles, corrections UX", "10 – 15 j"],
        ["Phase 3 — Vidéo investisseurs", "Scénarios de démo, visuels", "5 – 8 j"],
        ["Phase 4 — Version investisseurs", "Trading, monitoring IoT, multi-langues, Grands Comptes", "30 – 40 j"],
        ["Phase 5 — Production complète", "IA propriétaire, déploiement international, cybersécurité", "35 – 50 j"],
        ["TOTAL V1 complète", "", "~118 – 163 jours"],
    ],
    col_widths=[5.5, 7.5, 6.0]
)

p = doc.add_paragraph()
run1 = p.add_run("TJM de référence (senior / spécialité énergie + IA) : ")
set_font(run1, size=10.5, bold=True)
run2 = p.add_run("650 – 900 €/jour")
set_font(run2, size=10.5, bold=True, color=(20, 100, 20))
p.paragraph_format.space_after = Pt(10)

heading2(doc, "B. Trois modèles tarifaires proposés")

# Option A
p = doc.add_paragraph()
run = p.add_run("Option A — Forfait pur (classique prestataire)")
set_font(run, size=11, bold=True, color=(180, 80, 0))
p.paragraph_format.space_before = Pt(6)
bullet(doc, "Bêta + version investisseurs (phases 0→4) : 55 000 – 80 000 €")
bullet(doc, "V1 complète (toutes phases) : 90 000 – 140 000 €")
bullet(doc, "Avantage : sécurité financière immédiate. Inconvénient : vous cédez beaucoup de valeur.")

# Option B
p = doc.add_paragraph()
run = p.add_run("Option B — Apport réduit + royalties (co-éditeur)  ★ RECOMMANDÉE ★")
set_font(run, size=11, bold=True, color=(20, 100, 20))
p.paragraph_format.space_before = Pt(6)
bullet(doc, "Développement facturé à 40-50 % du forfait : 30 000 – 50 000 €")
bullet(doc, "Plus 15 à 25 % des revenus nets du logiciel (licences, abonnements SaaS)")
bullet(doc, (
    "Ce modèle est cohérent avec la proposition de co-édition et co-titularité PI. "
    "L'upside est potentiellement bien supérieur si le logiciel se déploie à l'international."
))

# Option C
p = doc.add_paragraph()
run = p.add_run("Option C — Mode SaaS abonnement (si hébergement géré en propre)")
set_font(run, size=11, bold=True, color=(80, 40, 140))
p.paragraph_format.space_before = Pt(6)
bullet(doc, "Développement initial : 25 000 – 40 000 €")
bullet(doc, "Plus abonnement annuel par licence : 4 000 – 12 000 €/an selon le tier (PME / Grands Comptes)")
bullet(doc, "Revenu récurrent. À coupler avec maintenance et mises à jour paramétriques.")

doc.add_paragraph()

heading2(doc, "C. Points de vigilance financiers")

bullet(doc, (
    "PI / levier de négociation : si NEWS-SOLAR apporte ses bases de données techniques et modèles de rendement, "
    "la valeur de la contribution logicielle (architecture + code) est équitable à 40–60 % de la valeur totale "
    "du logiciel."
))
bullet(doc, (
    "Minimum garanti : même en co-édition, prévoir un plancher de 20 000–30 000 € non remboursable "
    "pour couvrir le développement bêta."
))
bullet(doc, (
    "Protection du code : licence d'utilisation exclusive dans le domaine NEWS-SOLAR, conservation du droit "
    "de réutiliser les briques génériques sur d'autres projets."
))
bullet(doc, (
    "Budget maintenance annuel : 10–15 % du forfait initial, soit 9 000–15 000 €/an pour les mises à jour "
    "matérielles et réglementaires."
))

doc.add_paragraph()

# Recommandation encadrée
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Cm(1)
p.paragraph_format.right_indent = Cm(1)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run(
    "Recommandation : proposer l'Option B — 35 000 € de développement garanti (phases 0→4) + 20 % des "
    "revenus logiciel nets. Crédible, équitable, et aligné avec la logique de partenaire co-éditeur."
)
set_font(run, size=11, bold=True, italic=True, color=(20, 70, 140))

# Bordure sur le paragraphe de recommandation
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
for side in ("top", "bottom", "left", "right"):
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    "8")
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), "1F5C9E")
    pBdr.append(bdr)
pPr.append(pBdr)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE V — VENTILATION DES PAIEMENTS PAR JALONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "V. Ventilation des paiements par jalons — Option B retenue")

body(doc, (
    "Sur la base de l'Option B : 35 000 € forfait garanti + 20 % des revenus nets logiciel. "
    "Le forfait est structuré en 5 jalons d'avancement, chacun déclenché par une livraison vérifiable."
))

doc.add_paragraph()
heading2(doc, "A. Partie fixe : 35 000 € — mensualités + complément à livraison")

body(doc, (
    "Structure retenue : forfait mensuel de 3 800 € sur 6 mois (durée du projet phases 0→5), "
    "complété d'un solde de 12 200 € versé à la livraison et validation de la V1 complète. "
    "Total partie fixe : 6 × 3 800 € + 12 200 € = 35 000 €."
))

doc.add_paragraph()

# Sous-tableau 1 : mensualités
p = doc.add_paragraph()
run = p.add_run("Mensualités (facturation le 1er de chaque mois)")
set_font(run, size=10.5, bold=True, italic=True, color=(20, 70, 140))
p.paragraph_format.space_after = Pt(4)

add_table(
    doc,
    headers=["Mois", "Période indicative", "Phases actives", "Montant mensuel"],
    rows=[
        ["M1", "Mois 1",   "Phase 0 — Cadrage & contractualisation JDA",                        "3 800 €"],
        ["M2", "Mois 2",   "Phase 1 — Développement bêta (GPS / simulation / moteur HST)",       "3 800 €"],
        ["M3", "Mois 3",   "Phase 1 fin + Phase 2 — Ajustements & recalibrage modèles",          "3 800 €"],
        ["M4", "Mois 4",   "Phase 3 — Vidéo investisseurs + Phase 4 début (trading/IoT)",        "3 800 €"],
        ["M5", "Mois 5",   "Phase 4 fin — Multi-langues, Grands Comptes, démo investisseurs",    "3 800 €"],
        ["M6", "Mois 6",   "Phase 5 — Déploiement production, cybersécurité, formation",         "3 800 €"],
        ["",   "SOUS-TOTAL mensualités", "",                                                      "22 800 €"],
    ],
    col_widths=[1.2, 2.8, 8.0, 3.0]
)

doc.add_paragraph()

# Sous-tableau 2 : complément livraison
p = doc.add_paragraph()
run = p.add_run("Complément à la livraison V1 (versé après validation écrite)")
set_font(run, size=10.5, bold=True, italic=True, color=(20, 70, 140))
p.paragraph_format.space_after = Pt(4)

add_table(
    doc,
    headers=["Jalon", "Déclencheur de paiement", "Montant"],
    rows=[
        ["Livraison & validation V1",
         "V1 déployée en production + formation équipe NEWS-SOLAR effectuée + PV de recette signé",
         "12 200 €"],
        ["TOTAL PARTIE FIXE", "22 800 € mensualités + 12 200 € complément", "35 000 €"],
    ],
    col_widths=[3.8, 10.2, 3.0]
)

doc.add_paragraph()
heading2(doc, "B. Partie variable : 20 % des revenus nets logiciel")

add_table(
    doc,
    headers=["Mécanisme", "Détail"],
    rows=[
        ["Assiette",
         "Revenus nets perçus par NEWS-SOLAR liés au logiciel (licences, SaaS, modules additionnels) — hors TVA, hors éventuels revendeurs tiers"],
        ["Fréquence",
         "Reporting trimestriel avec versement sous 30 jours après clôture du trimestre"],
        ["Reporting",
         "Tableau de bord partagé + relevé signé par la direction NEWS-SOLAR"],
        ["Minimum annuel",
         "0 € la 1ère année (phase de lancement) — plancher de 5 000 €/an à partir de l'an 2"],
        ["Durée",
         "Sur toute la durée d'exploitation commerciale du logiciel, ou jusqu'à rachat de la quote-part"],
        ["Clause de rachat",
         "Possibilité pour NEWS-SOLAR de racheter la quote-part variable à tout moment sur la base d'un multiple de 3× les revenus annuels moyens des 2 dernières années"],
    ],
    col_widths=[4.0, 12.0]
)

doc.add_paragraph()
heading2(doc, "C. Points de protection contractuels")

bullet(doc, (
    "Mensualités non conditionnelles : les mensualités M1–M6 sont dues indépendamment de l'état "
    "d'avancement, dès lors que le prestataire est en activité sur le projet. Tout retard imputable "
    "à NEWS-SOLAR (absence de fourniture de données, validation tardive) n'en suspend pas le versement."
))
bullet(doc, (
    "Complément conditionnel : les 12 200 € sont versés uniquement à la signature du PV de recette "
    "V1. Délai de recette : 15 jours ouvrés après livraison — passé ce délai sans retour écrit motivé, "
    "la recette est réputée acceptée et le complément devient exigible."
))
bullet(doc, (
    "Gel du projet : si NEWS-SOLAR suspend le projet plus de 30 jours consécutifs, la mensualité "
    "du mois en cours reste due à 100 %. Au-delà de 60 jours de suspension, le complément de "
    "12 200 € devient exigible à 50 % au titre de réservation de capacité."
))
bullet(doc, (
    "Protection du code : licence d'utilisation exclusive dans le domaine NEWS-SOLAR, conservation "
    "du droit de réutiliser les briques génériques sur d'autres projets."
))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTIE VI — PROCHAINE ÉTAPE
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, "VI. Prochaine étape suggérée")

body(doc, (
    "Je propose que nous convenions d'une session de travail technique de 2–3 heures dans les meilleurs "
    "délais, avec les objectifs suivants :"
))
bullet(doc, "Revue de votre base de données technique (modèles de rendement, typologies d'installation, données IA)")
bullet(doc, "Démonstration live de la plateforme actuelle et identification des briques directement réutilisables")
bullet(doc, "Accord sur le périmètre exact de la bêta et le calendrier Phase 0/1")
bullet(doc, "Premiers éléments de la structure contractuelle JDA")

doc.add_paragraph()
body(doc, "Dans l'attente de votre retour, je reste disponible pour convenir d'une date au plus tôt.")
doc.add_paragraph()
body(doc, "Cordialement,")

# ── Pied de page ────────────────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Document confidentiel — Proposition de collaboration NEWS-SOLAR / Logiciel HST")
    set_font(run, size=8, italic=True, color=(150, 150, 150))

# ── Sauvegarde ───────────────────────────────────────────────────────────────
out_path = "NEWS_SOLAR_Proposition_Collaboration.docx"
doc.save(out_path)
print(f"Document généré : {out_path}")
