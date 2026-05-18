"""
Génération du CV de Yann Laurent en format Word (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_FILE = "CV_Yann_Laurent.docx"

# ── Couleurs ──────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x37, 0x5E)   # titres principaux
MID_BLUE   = RGBColor(0x1F, 0x6F, 0xAA)   # section headers
ORANGE     = RGBColor(0xE8, 0x7A, 0x1E)   # accents / bullet
GREY_TEXT  = RGBColor(0x44, 0x44, 0x44)   # corps de texte
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)


def set_run_font(run, name="Calibri", size_pt=10, bold=False, italic=False,
                 color=None):
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_horizontal_rule(doc, color_hex="1F6FAA", thickness=12):
    """Insère un filet horizontal coloré sous le dernier paragraphe."""
    p = doc.paragraphs[-1]._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_paragraph(paragraph, fill_hex="1A375E"):
    """Colorie le fond d'un paragraphe."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line)


def add_section_title(doc, text):
    """Titre de section avec fond bleu foncé et texte blanc."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=2)
    p.paragraph_format.left_indent = Cm(0)
    shade_paragraph(p, "1A375E")
    run = p.add_run(f"  {text.upper()}")
    set_run_font(run, size_pt=11, bold=True, color=WHITE)
    return p


def add_job(doc, title, company, location_period, bullets):
    """Bloc expérience professionnelle."""
    # Ligne poste / entreprise
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=0)
    r_title = p.add_run(title)
    set_run_font(r_title, size_pt=10.5, bold=True, color=DARK_BLUE)
    r_sep = p.add_run("  –  ")
    set_run_font(r_sep, size_pt=10.5, color=GREY_TEXT)
    r_company = p.add_run(company)
    set_run_font(r_company, size_pt=10.5, bold=True, color=MID_BLUE)

    # Lieu / période
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=2)
    r_loc = p2.add_run(location_period)
    set_run_font(r_loc, size_pt=9, italic=True, color=GREY_TEXT)

    # Bullets
    for b in bullets:
        p3 = doc.add_paragraph(style="List Bullet")
        p3.paragraph_format.left_indent  = Cm(0.6)
        p3.paragraph_format.space_before = Pt(1)
        p3.paragraph_format.space_after  = Pt(1)
        run = p3.add_run(b)
        set_run_font(run, size_pt=9.5, color=GREY_TEXT)


def add_competence_item(doc, label, detail):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(label + " : ")
    set_run_font(r1, size_pt=9.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run(detail)
    set_run_font(r2, size_pt=9.5, color=GREY_TEXT)


# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Marges
section = doc.sections[0]
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)

# ── EN-TÊTE NOM ───────────────────────────────────────────────────────────────
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_name, before=0, after=2)
r = p_name.add_run("YANN LAURENT")
set_run_font(r, size_pt=24, bold=True, color=DARK_BLUE)

# Ligne de contact
p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_contact, before=0, after=2)
contact_text = (
    "13 Ventenat – 23480 Saint Sulpice les Champs, France\n"
    "+33 (0) 6 21 16 55 85  |  eco.enr@gmail.com\n"
    "Français natif  |  Anglais courant  |  Italien & Espagnol opérationnels"
)
r_c = p_contact.add_run(contact_text)
set_run_font(r_c, size_pt=9, color=GREY_TEXT)

# Filet sous l'en-tête
add_horizontal_rule(doc, color_hex="E87A1E", thickness=16)

# Titre de poste
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_title, before=6, after=4)
r_t = p_title.add_run(
    "Directeur Commercial & Développement  –  Expert Industrie Photovoltaïque"
)
set_run_font(r_t, size_pt=13, bold=True, color=MID_BLUE)

# Résumé profil
p_summary = doc.add_paragraph()
set_paragraph_spacing(p_summary, before=2, after=6)
p_summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
summary = (
    "20 ans d'expérience dans l'énergie solaire | Expertise 360° de la chaîne de valeur photovoltaïque\n"
    "Leader stratégique et opérationnel dans les énergies renouvelables, maîtrisant l'ensemble des segments "
    "du photovoltaïque : EPC, développement de projets, innovation produit, autoconsommation, agrivoltaïsme "
    "et centrales au sol de grande envergure. Reconnu pour sa capacité à structurer des équipes, construire "
    "des processus et délivrer des résultats à fort impact en Europe, au Moyen-Orient et en Afrique."
)
r_s = p_summary.add_run(summary)
set_run_font(r_s, size_pt=9.5, color=GREY_TEXT)

# ── EXPÉRIENCES PROFESSIONNELLES ──────────────────────────────────────────────
add_section_title(doc, "Expériences Professionnelles")

EXPERIENCES = [
    (
        "Directeur Pays", "SUN SUPPORT",
        "France  |  2023 – 2024",
        [
            "Constitution et formation de nouvelles équipes commerciales et techniques",
            "Refonte de la communication corporate et de l'offre produits",
            "Lancement de nouvelles solutions auprès de clients majeurs : TotalEnergies, Elements, Photosol, Engie",
            "Gestion des comptes stratégiques et des partenariats clés",
        ],
    ),
    (
        "Directeur Développement", "PHENAE SOLUTION",
        "France  |  2022 – 2023",
        [
            "Création et pilotage des équipes commerciales et bureau d'études",
            "Élaboration des supports commerciaux et formation des apporteurs d'affaires",
            "Réalisation de 2,2 M€ de chiffre d'affaires la première année – pipeline > 10 M€ en 2023",
        ],
    ),
    (
        "Directeur Développement", "RECOM SOLAR / RECOM ENERGY",
        "Europe  |  2019 – 2021",
        [
            "Vente de modules PV et de portefeuilles de projets à des acteurs majeurs (EDF EN, ENGIE, NEOEN)",
            "Pilotage d'un appel d'offres 200 MWp de bout en bout, ayant attiré plus de 15 soumissionnaires "
            "(Shell, ENEL, Statkraft, Photosol…)",
            "Conduite de transactions sur des projets prêts à construire (RTB)",
        ],
    ),
    (
        "Directeur Développement", "SUN'R",
        "France  |  2017 – 2019",
        [
            "Gestion du flux de développement de projets PV de grande envergure (toitures, centrales au sol, agrivoltaïsme)",
            "Accompagnement des équipes projet dans l'identification et la maturation des opportunités viables",
        ],
    ),
    (
        "Business Developer International", "BITRON SPA",
        "MENA, Afrique de l'Ouest, Europe  |  2012 – 2016",
        [
            "Développement d'une solution d'optimisation PV (hardware/software)",
            "Suivi de projets de taille intermédiaire (2 – 50 MWp) avec des entités gouvernementales "
            "(ex. Ministère de l'Énergie du Maroc)",
            "Négociation de PPA, coordination des activités EPC, études de faisabilité",
        ],
    ),
    (
        "Chef de Projets / Ingénieur Commercial", "THERMOVOLT AG",
        "Suisse  |  2009 – 2012",
        [
            "Réalisation EPC de projets PV de 100 kWp à 1 MWp",
            "Gestion de la construction et de l'ingénierie de plus de 4 MWp de systèmes en toiture",
        ],
    ),
    (
        "Directeur France", "SYSTEM S.p.A",
        "France  |  2009 – 2011",
        [
            "Lancement d'une nouvelle tuile photovoltaïque BIPV sur le marché français",
            "Constitution d'un réseau national de distributeurs ; introduction produit chez Leroy Merlin / Auchan",
            "Conception de la stratégie go-to-market complète (marketing, vente, support)",
        ],
    ),
    (
        "Ingénieur Technico-Commercial", "ECOSTREAM France",
        "France  |  2006 – 2008",
        [
            "Développement d'un réseau de distribution B2B pour kits solaires (1 – 3 kWp)",
            "Création d'outils d'aide à la vente (Solartool), gestion des offres EPC et négociations contractuelles",
        ],
    ),
]

for title, company, loc_period, bullets in EXPERIENCES:
    add_job(doc, title, company, loc_period, bullets)

# ── COMPÉTENCES CLÉS ──────────────────────────────────────────────────────────
add_section_title(doc, "Compétences Clés")

COMPETENCES = [
    ("Expertise PV", "composants, développement de projets, EPC, BIPV, agrivoltaïsme, autoconsommation"),
    ("Développement stratégique", "structuration d'équipes, développement commercial, lancement produit, gestion grands comptes"),
    ("Gestion de projets", "de la faisabilité à l'exécution : appels d'offres, négociation de contrats"),
    ("Innovation", "stratégie produit, solutions techniques, repositionnement marketing"),
    ("Business international", "réseau étendu Europe, MENA & Afrique"),
    ("Programmation Python", "traitement de données, automatisation, applications énergétiques"),
]

for label, detail in COMPETENCES:
    add_competence_item(doc, label, detail)

# ── OUTILS & LOGICIELS ────────────────────────────────────────────────────────
add_section_title(doc, "Outils & Logiciels")

p_tools = doc.add_paragraph()
set_paragraph_spacing(p_tools, before=4, after=2)
p_tools.paragraph_format.left_indent = Cm(0.6)
r_tools = p_tools.add_run(
    "Suite Microsoft Office  |  Adobe InDesign  |  Adobe Photoshop  |  "
    "PVsyst  |  PVSol  |  SketchUp  |  Python"
)
set_run_font(r_tools, size_pt=9.5, color=GREY_TEXT)

# ── LANGUES ───────────────────────────────────────────────────────────────────
add_section_title(doc, "Langues")

p_lang = doc.add_paragraph()
set_paragraph_spacing(p_lang, before=4, after=2)
p_lang.paragraph_format.left_indent = Cm(0.6)
languages = [
    ("Français", "Langue maternelle"),
    ("Anglais", "Courant"),
    ("Italien", "Opérationnel"),
    ("Espagnol", "Opérationnel"),
]
for i, (lang, level) in enumerate(languages):
    r1 = p_lang.add_run(lang + " : ")
    set_run_font(r1, size_pt=9.5, bold=True, color=DARK_BLUE)
    r2 = p_lang.add_run(level)
    set_run_font(r2, size_pt=9.5, color=GREY_TEXT)
    if i < len(languages) - 1:
        r_sep = p_lang.add_run("   |   ")
        set_run_font(r_sep, size_pt=9.5, color=GREY_TEXT)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
doc.save(OUTPUT_FILE)
print(f"CV généré avec succès : {OUTPUT_FILE}")
