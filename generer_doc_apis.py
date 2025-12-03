#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Génération d'un document Word détaillant les API utilisées par AgriWeb
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Ajoute un lien hypertexte à un paragraphe"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style du lien
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def create_api_documentation():
    """Crée le document Word complet"""
    doc = Document()
    
    # === STYLES ===
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # === PAGE DE GARDE ===
    title = doc.add_heading('AgriWeb - Documentation des API', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Liste intégrale des couches et endpoints utilisés')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(68, 114, 196)
    
    doc.add_paragraph()
    
    date_p = doc.add_paragraph('Document généré le 21 novembre 2025')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_format = date_p.runs[0]
    date_format.font.size = Pt(10)
    date_format.font.italic = True
    date_format.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # === TABLE DES MATIÈRES (manuelle) ===
    doc.add_heading('Table des matières', 1)
    toc_items = [
        '1. API Cadastre IGN (APIcarto Cadastre)',
        '2. API Nature IGN (APIcarto Nature)',
        '3. API GPU Urbanisme IGN (APIcarto GPU)',
        '4. API GéoRisques',
        '5. API Géographiques',
        '6. API RPG (Registre Parcellaire Graphique)',
        '7. API Altitude IGN',
        '8. Résumé statistique',
        '9. Informations complémentaires'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # === 1. API CADASTRE ===
    doc.add_heading('1. API Cadastre IGN (APIcarto Cadastre)', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API Cadastre permet de récupérer les références cadastrales des parcelles "
        "à partir de coordonnées géographiques ou de géométries."
    )
    
    doc.add_heading('URL de base', 2)
    p = doc.add_paragraph()
    p.add_run('https://apicarto.ign.fr/api/cadastre').bold = True
    
    doc.add_heading('Endpoints utilisés', 2)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # En-têtes
    headers = table.rows[0].cells
    headers[0].text = 'Endpoint'
    headers[1].text = 'Description'
    headers[2].text = 'Usage dans AgriWeb'
    
    for cell in headers:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4472C4')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Données
    row = table.rows[1].cells
    row[0].text = '/parcelle'
    row[1].text = 'Récupération des parcelles cadastrales'
    row[2].text = 'Enrichissement parkings, toitures, friches, parcelles RPG'
    
    doc.add_heading('Paramètres utilisés', 2)
    params = doc.add_paragraph()
    params.add_run('• geom').bold = True
    params.add_run(' : GeoJSON Point ou Polygon\n')
    params.add_run('• _limit').bold = True
    params.add_run(' : 1000 (nombre maximum de résultats)\n')
    params.add_run('• source_ign').bold = True
    params.add_run(' : PCI (Plan Cadastral Informatisé)')
    
    doc.add_heading('Exemple de requête', 2)
    code = doc.add_paragraph(
        'GET https://apicarto.ign.fr/api/cadastre/parcelle?'
        'geom={"type":"Point","coordinates":[2.3522,48.8566]}&_limit=1000&source_ign=PCI',
        style='Intense Quote'
    )
    
    doc.add_page_break()
    
    # === 2. API NATURE ===
    doc.add_heading('2. API Nature IGN (APIcarto Nature)', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API Nature fournit l'accès aux zones naturelles protégées : "
        "Natura 2000, ZNIEFF, Parcs Nationaux et Régionaux, Réserves Naturelles."
    )
    
    doc.add_heading('URL de base', 2)
    p = doc.add_paragraph()
    p.add_run('https://apicarto.ign.fr/api/nature').bold = True
    
    doc.add_heading('Endpoints utilisés (9 couches)', 2)
    
    # Tableau des endpoints nature
    nature_table = doc.add_table(rows=10, cols=3)
    nature_table.style = 'Light Grid Accent 1'
    
    # En-têtes
    nature_headers = nature_table.rows[0].cells
    nature_headers[0].text = 'Endpoint'
    nature_headers[1].text = 'Nom complet'
    nature_headers[2].text = 'Type de protection'
    
    for cell in nature_headers:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '70AD47')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Données
    nature_data = [
        ('/natura-habitat', 'Natura 2000 Directive Habitat', 'Européenne'),
        ('/natura-oiseaux', 'Natura 2000 Directive Oiseaux', 'Européenne'),
        ('/znieff1', 'ZNIEFF Type 1', 'Nationale'),
        ('/znieff2', 'ZNIEFF Type 2', 'Nationale'),
        ('/pn', 'Parcs Nationaux', 'Nationale'),
        ('/pnr', 'Parcs Naturels Régionaux', 'Régionale'),
        ('/rnn', 'Réserves Naturelles Nationales', 'Nationale'),
        ('/rnc', 'Réserves Naturelles de Corse', 'Régionale'),
        ('/rncf', 'Réserves Nationales Chasse Faune Sauvage', 'Nationale'),
    ]
    
    for i, (endpoint, nom, type_prot) in enumerate(nature_data, 1):
        cells = nature_table.rows[i].cells
        cells[0].text = endpoint
        cells[1].text = nom
        cells[2].text = type_prot
    
    doc.add_heading('Paramètres utilisés', 2)
    params = doc.add_paragraph()
    params.add_run('• geom').bold = True
    params.add_run(' : GeoJSON Polygon (contour de la commune)\n')
    params.add_run('• _limit').bold = True
    params.add_run(' : 1000')
    
    doc.add_heading('Usage', 2)
    doc.add_paragraph(
        "Ces données sont utilisées pour :\n"
        "• Identifier les contraintes environnementales\n"
        "• Afficher les zones protégées sur la carte interactive\n"
        "• Générer les alertes dans le rapport de faisabilité"
    )
    
    doc.add_page_break()
    
    # === 3. API GPU URBANISME ===
    doc.add_heading('3. API GPU Urbanisme IGN (APIcarto GPU)', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API GPU (Géoportail de l'Urbanisme) donne accès aux documents d'urbanisme "
        "et aux servitudes d'utilité publique."
    )
    
    doc.add_heading('URL de base', 2)
    p = doc.add_paragraph()
    p.add_run('https://apicarto.ign.fr/api/gpu').bold = True
    
    doc.add_heading('Endpoints utilisés (17 couches)', 2)
    
    # Tableau GPU
    gpu_table = doc.add_table(rows=18, cols=3)
    gpu_table.style = 'Light Grid Accent 1'
    
    # En-têtes
    gpu_headers = gpu_table.rows[0].cells
    gpu_headers[0].text = 'Endpoint'
    gpu_headers[1].text = 'Description'
    gpu_headers[2].text = 'Catégorie'
    
    for cell in gpu_headers:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FFC000')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    gpu_data = [
        ('municipality', 'Communes', 'Administratif'),
        ('document', 'Documents d\'urbanisme', 'Administratif'),
        ('zone-urba', 'Zones d\'urbanisme (PLU/POS)', 'Zonage'),
        ('secteur-cc', 'Secteurs de coefficient constructibilité', 'Zonage'),
        ('prescription-surf', 'Prescriptions surfaciques', 'Prescription'),
        ('prescription-lin', 'Prescriptions linéaires', 'Prescription'),
        ('prescription-pct', 'Prescriptions ponctuelles', 'Prescription'),
        ('info-surf', 'Informations surfaciques', 'Information'),
        ('info-lin', 'Informations linéaires', 'Information'),
        ('info-pct', 'Informations ponctuelles', 'Information'),
        ('acte-sup', 'Actes SUP', 'SUP'),
        ('assiette-sup-s', 'Assiettes SUP surfaciques', 'SUP'),
        ('assiette-sup-l', 'Assiettes SUP linéaires', 'SUP'),
        ('assiette-sup-p', 'Assiettes SUP ponctuelles', 'SUP'),
        ('generateur-sup-s', 'Générateurs SUP surfaciques', 'SUP'),
        ('generateur-sup-l', 'Générateurs SUP linéaires', 'SUP'),
        ('generateur-sup-p', 'Générateurs SUP ponctuels', 'SUP'),
    ]
    
    for i, (endpoint, desc, cat) in enumerate(gpu_data, 1):
        cells = gpu_table.rows[i].cells
        cells[0].text = endpoint
        cells[1].text = desc
        cells[2].text = cat
    
    doc.add_heading('Endpoint principal : zone-urba', 2)
    doc.add_paragraph(
        "L'endpoint zone-urba est le plus utilisé car il fournit le zonage PLU/POS "
        "nécessaire pour déterminer la constructibilité (zones A, N, U, AU, etc.)."
    )
    
    doc.add_page_break()
    
    # === 4. API GEORISQUES ===
    doc.add_heading('4. API GéoRisques', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API GéoRisques du Ministère de la Transition Écologique fournit les données "
        "sur les risques naturels et technologiques."
    )
    
    doc.add_heading('URL de base', 2)
    p = doc.add_paragraph()
    p.add_run('https://www.georisques.gouv.fr/api/v1').bold = True
    
    doc.add_heading('Endpoints utilisés (9 couches)', 2)
    
    georisques_table = doc.add_table(rows=10, cols=2)
    georisques_table.style = 'Light Grid Accent 1'
    
    georisques_headers = georisques_table.rows[0].cells
    georisques_headers[0].text = 'Endpoint'
    georisques_headers[1].text = 'Description'
    
    for cell in georisques_headers:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C00000')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    georisques_data = [
        ('/zonage_sismique', 'Zonage sismique'),
        ('/gaspar/catnat', 'Catastrophes naturelles'),
        ('/cavites', 'Cavités souterraines'),
        ('/mvt', 'Mouvements de terrain'),
        ('/gaspar/risques', 'Risques GASPAR'),
        ('/radon', 'Potentiel radon'),
        ('/installations_classees', 'ICPE (Installations Classées)'),
        ('/installations_nucleaires', 'Installations nucléaires'),
        ('/zonage/pprn', '⭐ PPRI/PPRN (Plans Prévention Risques)'),
    ]
    
    for i, (endpoint, desc) in enumerate(georisques_data, 1):
        cells = georisques_table.rows[i].cells
        cells[0].text = endpoint
        cells[1].text = desc
    
    doc.add_heading('Focus : Endpoint PPRI/PPRN', 2)
    doc.add_paragraph(
        "L'endpoint /zonage/pprn est particulièrement important pour identifier les zones "
        "inondables et autres risques naturels prévisibles."
    )
    
    params = doc.add_paragraph()
    params.add_run('Paramètres PPRI :\n').bold = True
    params.add_run('• lat').bold = True
    params.add_run(' : Latitude du point\n')
    params.add_run('• lon').bold = True
    params.add_run(' : Longitude du point\n')
    params.add_run('• rayon').bold = True
    params.add_run(' : Rayon de recherche en mètres (défaut : 1000m)\n')
    params.add_run('• format').bold = True
    params.add_run(' : geojson')
    
    doc.add_page_break()
    
    # === 5. API GEOGRAPHIQUES ===
    doc.add_heading('5. API Géographiques (geo.api.gouv.fr)', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API Géo fournit les données administratives : communes, départements, régions."
    )
    
    doc.add_heading('URL de base', 2)
    p = doc.add_paragraph()
    p.add_run('https://geo.api.gouv.fr').bold = True
    
    doc.add_heading('Endpoints utilisés', 2)
    
    geo_table = doc.add_table(rows=4, cols=2)
    geo_table.style = 'Light Grid Accent 1'
    
    geo_headers = geo_table.rows[0].cells
    geo_headers[0].text = 'Endpoint'
    geo_headers[1].text = 'Usage'
    
    for cell in geo_headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    geo_data = [
        ('/communes?nom={nom}&fields=centre,contour,code,population,surface', 
         'Géocodage et contour de commune'),
        ('/communes?nom={nom}&fields=mairie', 
         'Coordonnées de la mairie'),
        ('/departements/{code}/communes', 
         'Liste des communes (autocomplete)'),
    ]
    
    for i, (endpoint, usage) in enumerate(geo_data, 1):
        cells = geo_table.rows[i].cells
        cells[0].text = endpoint
        cells[1].text = usage
    
    # === 6. API RPG ===
    doc.add_heading('6. API RPG (Registre Parcellaire Graphique)', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "L'API RPG fournit les données du Registre Parcellaire Graphique : "
        "parcelles agricoles déclarées à la PAC avec leurs cultures."
    )
    
    doc.add_heading('URL', 2)
    p = doc.add_paragraph()
    p.add_run('https://apicarto.ign.fr/api/rpg/parcelles').bold = True
    
    doc.add_heading('Paramètres', 2)
    params = doc.add_paragraph()
    params.add_run('• geom').bold = True
    params.add_run(' : GeoJSON Polygon (contour commune)')
    
    # === 7. API ALTITUDE ===
    doc.add_heading('7. API Altitude IGN', 1)
    
    doc.add_heading('Description', 2)
    doc.add_paragraph(
        "Service de calcul d'altitude à partir de coordonnées géographiques."
    )
    
    doc.add_heading('URL', 2)
    p = doc.add_paragraph()
    p.add_run('https://wxs.ign.fr/calcul/alti/rest/elevation.json').bold = True
    
    doc.add_heading('Paramètres', 2)
    params = doc.add_paragraph()
    params.add_run('• lon').bold = True
    params.add_run(' : Longitude\n')
    params.add_run('• lat').bold = True
    params.add_run(' : Latitude')
    
    doc.add_page_break()
    
    # === 8. RÉSUMÉ ===
    doc.add_heading('8. Résumé statistique', 1)
    
    resume_table = doc.add_table(rows=9, cols=3)
    resume_table.style = 'Medium Shading 1 Accent 1'
    
    resume_headers = resume_table.rows[0].cells
    resume_headers[0].text = 'Catégorie'
    resume_headers[1].text = 'Nombre d\'endpoints'
    resume_headers[2].text = 'API'
    
    for cell in resume_headers:
        cell.paragraphs[0].runs[0].font.bold = True
    
    resume_data = [
        ('Cadastre', '1', 'APIcarto IGN'),
        ('Nature', '9', 'APIcarto IGN'),
        ('Urbanisme (GPU)', '17', 'APIcarto IGN'),
        ('Risques (PPRI)', '9', 'GéoRisques'),
        ('Géographie', '3', 'API Géo Gouv'),
        ('Agriculture (RPG)', '1', 'APIcarto IGN'),
        ('Altitude', '1', 'IGN Géoportail'),
        ('TOTAL', '41 endpoints', ''),
    ]
    
    for i, (cat, nb, api) in enumerate(resume_data, 1):
        cells = resume_table.rows[i].cells
        cells[0].text = cat
        cells[1].text = nb
        cells[2].text = api
        if cat == 'TOTAL':
            for cell in cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # === 9. INFORMATIONS COMPLÉMENTAIRES ===
    doc.add_heading('9. Informations complémentaires', 1)
    
    doc.add_heading('Documentation officielle', 2)
    p = doc.add_paragraph()
    p.add_run('• APIcarto : ').bold = True
    add_hyperlink(p, 'https://apicarto.ign.fr/api/doc/', 'https://apicarto.ign.fr/api/doc/')
    
    p = doc.add_paragraph()
    p.add_run('• GéoRisques : ').bold = True
    add_hyperlink(p, 'https://www.georisques.gouv.fr/doc-api', 'https://www.georisques.gouv.fr/doc-api')
    
    p = doc.add_paragraph()
    p.add_run('• API Géo : ').bold = True
    add_hyperlink(p, 'https://geo.api.gouv.fr/', 'https://geo.api.gouv.fr/')
    
    doc.add_heading('Limites de requêtes', 2)
    limits = doc.add_paragraph()
    limits.add_run('• Limite par défaut : ').bold = True
    limits.add_run('_limit=1000 pour APIcarto\n')
    limits.add_run('• Clé API : ').bold = True
    limits.add_run('Non nécessaire (APIs publiques)\n')
    limits.add_run('• Timeout : ').bold = True
    limits.add_run('10 secondes par requête\n')
    limits.add_run('• Rate limiting : ').bold = True
    limits.add_run('Non spécifié (usage raisonnable recommandé)')
    
    doc.add_heading('Format de données', 2)
    format_p = doc.add_paragraph()
    format_p.add_run('• Format retourné : ').bold = True
    format_p.add_run('GeoJSON (principalement)\n')
    format_p.add_run('• Système de coordonnées : ').bold = True
    format_p.add_run('WGS84 (EPSG:4326)\n')
    format_p.add_run('• Encodage : ').bold = True
    format_p.add_run('UTF-8')
    
    doc.add_heading('Usage dans AgriWeb', 2)
    doc.add_paragraph(
        "Ces API sont interrogées de manière séquentielle lors de la génération "
        "d'un rapport de faisabilité pour une commune ou une adresse. "
        "Les données sont ensuite agrégées, analysées et affichées sur une carte "
        "interactive Folium."
    )
    
    doc.add_heading('Architecture technique', 2)
    archi = doc.add_paragraph()
    archi.add_run('• Méthode HTTP : ').bold = True
    archi.add_run('GET\n')
    archi.add_run('• Bibliothèque Python : ').bold = True
    archi.add_run('requests\n')
    archi.add_run('• Gestion des erreurs : ').bold = True
    archi.add_run('try/except avec timeout\n')
    archi.add_run('• Cache : ').bold = True
    archi.add_run('Non implémenté (interrogation à chaque requête)\n')
    archi.add_run('• Parallélisation : ').bold = True
    archi.add_run('Non (requêtes séquentielles)')
    
    # === PIED DE PAGE ===
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Document généré automatiquement par AgriWeb')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.font.size = Pt(9)
    footer_format.font.italic = True
    footer_format.font.color.rgb = RGBColor(128, 128, 128)
    
    # === SAUVEGARDE ===
    filename = 'AgriWeb_Documentation_API_Complete.docx'
    doc.save(filename)
    print(f"✅ Document créé : {filename}")
    return filename

if __name__ == "__main__":
    create_api_documentation()
